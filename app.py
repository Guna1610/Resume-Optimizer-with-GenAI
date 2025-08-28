import os
import io
import re
import tempfile
import streamlit as st
from docx import Document
from dotenv import load_dotenv

# Import your existing engine
import final_resume_optimizer as ro

load_dotenv()

st.set_page_config(page_title="Resume Optimizer (GenAI)", page_icon="📝", layout="wide")

# ---------- Tiny helpers ----------
SECTION_NAMES = ["SUMMARY", "SKILLS", "PROJECT EXPERIENCE"]

def split_sections_plain(text: str):
    """
    Very simple splitter: returns dict {section_name: text_block}
    Looks for exact headings in SECTION_NAMES; captures until the next known heading (or end).
    """
    lines = text.splitlines()
    idxs = []
    for i, ln in enumerate(lines):
        ln_norm = " ".join(ln.strip().split()).upper()
        if ln_norm in SECTION_NAMES:
            idxs.append((ln_norm, i))
    idxs.sort(key=lambda x: x[1])

    out = {name: "" for name in SECTION_NAMES}
    for k in range(len(idxs)):
        sec, start = idxs[k]
        end = idxs[k + 1][1] if k + 1 < len(idxs) else len(lines)
        block = "\n".join(lines[start + 1:end]).strip("\n")
        out[sec] = block
    return out

def apply_overrides_for_prompt(full_resume_text: str, overrides: dict) -> str:
    """
    Build a temporary 'effective' resume text ONLY for the LLM prompt:
    - Replace sections present in overrides (text blocks).
    - This does NOT touch the .docx; we still use the original for formatting.
    """
    # naive: rebuild by scanning and replacing between headings
    lines = full_resume_text.splitlines()
    # first, find headings with their line idx
    marks = []
    for i, ln in enumerate(lines):
        if " ".join(ln.strip().split()).upper() in SECTION_NAMES:
            marks.append(i)
    if not marks:
        # no headings found; best effort — just append edited sections at the end
        parts = [full_resume_text]
        for sec in SECTION_NAMES:
            if overrides.get(sec):
                parts.append(f"\n{sec}\n{overrides[sec].strip()}")
        return "\n".join(parts)

    # rebuild with replacements
    new_lines = []
    i = 0
    while i < len(lines):
        ln = lines[i]
        name = " ".join(ln.strip().split()).upper()
        if name in SECTION_NAMES:
            # write heading
            new_lines.append(ln)
            # skip old block
            j = i + 1
            while j < len(lines):
                test = " ".join(lines[j].strip().split()).upper()
                if test in SECTION_NAMES:
                    break
                j += 1
            # insert override if present; else keep original block
            if overrides.get(name):
                new_lines.extend(overrides[name].splitlines())
            else:
                new_lines.extend(lines[i + 1:j])
            i = j
        else:
            new_lines.append(ln)
            i += 1
    return "\n".join(new_lines)

def save_uploaded_file(tmpdir, uf, name):
    path = os.path.join(tmpdir, name)
    with open(path, "wb") as f:
        f.write(uf.read())
    return path

def write_text_file(tmpdir, name, text):
    path = os.path.join(tmpdir, name)
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path

def rough_keyword_match_score(jd: str, resume_text: str) -> float:
    import re
    tok = lambda s: set(t for t in re.findall(r"[a-zA-Z0-9+#\.]+", s.lower()) if len(t) > 2)
    jset = tok(jd)
    rset = tok(resume_text)
    if not jset:
        return 0.0
    return 100.0 * len(jset & rset) / len(jset)

# ---------- UI: layout ----------
st.title("📝 Resume Optimizer (GenAI)")

# Global CSS tweaks (make JD text area look like a terminal)
st.markdown("""
<style>
/* Monospace, dark for the JD "terminal" */
textarea[aria-label="JD Terminal"] {
  background-color: #0e1117 !important;
  color: #e6e6e6 !important;
  font-family: ui-monospace, SFMono-Regular, Menlo, Consolas, "Liberation Mono", monospace !important;
  border-radius: 10px !important;
}
/* Make textareas a tad denser */
.stTextArea textarea { line-height: 1.35; }
</style>
""", unsafe_allow_html=True)

left, right = st.columns([0.57, 0.43], gap="large")

with left:
    st.subheader("1) Upload & Edit Resume (left)")

    resume_file = st.file_uploader("Resume (.docx)", type=["docx"], key="resume_up")
    resume_text_extracted = ""
    summary_edit = ""
    skills_edit = ""
    projects_edit = ""

    if resume_file:
        with tempfile.TemporaryDirectory() as tmp:
            temp_resume_path = save_uploaded_file(tmp, resume_file, "resume.docx")
            # Use your existing extractor
            resume_text_extracted = ro.extract_resume_text(temp_resume_path)

        st.markdown("**Preview (plain text)**")
        st.text_area("Resume Preview", resume_text_extracted, height=220, key="resume_preview", help="Plain text preview extracted from your .docx")

        # Quick section editors
        st.markdown("**Quick Edit — Only these sections (optional)**")
        secs = split_sections_plain(resume_text_extracted)
        summary_edit = st.text_area("SUMMARY (edit before optimize)", secs.get("SUMMARY", ""), height=110)
        skills_edit  = st.text_area("SKILLS (edit before optimize)", secs.get("SKILLS", ""), height=140,
                                    help="Keep bullets; the optimizer will still rewrite SKILLS to match the JD.")
        projects_edit = st.text_area("PROJECT EXPERIENCE (edit before optimize)", secs.get("PROJECT EXPERIENCE", ""), height=220)

    else:
        st.info("Upload a `.docx` resume to preview and edit its sections here.")

with right:
    st.subheader("2) Job Description Terminal (right)")

    jd_col1, jd_col2 = st.columns(2)
    with jd_col1:
        jd_file = st.file_uploader("JD (.txt, optional)", type=["txt"], key="jd_up")
    with jd_col2:
        projects_file = st.file_uploader("Project Library (.txt)", type=["txt"], key="proj_up",
                                         help="All your projects; the app will pick the top 3 that best match the JD.")

    jd_text = st.text_area("JD Terminal", height=260, key="jd_terminal",
                           placeholder="Paste the full job description here…")

    st.markdown("---")
    run_btn = st.button("🚀 Optimize Resume")

# ---------- Run pipeline ----------
if run_btn:
    if not resume_file:
        st.error("Please upload your resume (.docx).")
        st.stop()

    if not projects_file:
        st.error("Please upload your full project library (.txt).")
        st.stop()

    # Determine JD source
    if jd_file and not jd_text.strip():
        with tempfile.TemporaryDirectory() as tmp:
            jd_path = save_uploaded_file(tmp, jd_file, "jd.txt")
            with open(jd_path, "r", encoding="utf-8") as jf:
                jd_text = jf.read()
    elif not jd_text.strip() and not jd_file:
        st.error("Please paste the JD in the terminal or upload a JD file.")
        st.stop()

    try:
        with tempfile.TemporaryDirectory() as tmp:
            # Save original resume + project library
            resume_path = save_uploaded_file(tmp, resume_file, "resume.docx")
            proj_path = save_uploaded_file(tmp, projects_file, "projects.txt")

            # Build an 'effective' resume text for the prompt, honoring left-pane edits
            original_resume_text = ro.extract_resume_text(resume_path)
            overrides = {}
            if summary_edit.strip():
                overrides["SUMMARY"] = summary_edit.strip()
            if skills_edit.strip():
                overrides["SKILLS"] = skills_edit.strip()
            if projects_edit.strip():
                overrides["PROJECT EXPERIENCE"] = projects_edit.strip()

            effective_resume_text = apply_overrides_for_prompt(original_resume_text, overrides)

            # Use your existing Gemini call directly with the edited effective text
            projects_library_text = ro.extract_project_library(proj_path)
            optimized = ro.optimize_resume_sections(
                resume_text=effective_resume_text,
                job_text=jd_text,
                project_library=projects_library_text
            )

            # Apply into the original .docx (keeps design/format)
            doc = Document(resume_path)
            # If you don't want SUMMARY at all, comment the next line:
            ro.replace_section(doc, "SUMMARY", optimized["summary"])
            ro.replace_section(doc, "SKILLS", optimized["skills"])
            ro.replace_section(doc, "PROJECT EXPERIENCE", optimized["projects"])

            out_path = os.path.join(tmp, "optimized_resume.docx")
            doc.save(out_path)

            # Prepare download + small insights
            with open(out_path, "rb") as f:
                data = f.read()

            # Quick insights
            text_all = "\n".join(p.text for p in Document(out_path).paragraphs)
            score = rough_keyword_match_score(jd_text, text_all)

            st.success("Optimization complete! 🎉")
            st.write(f"**Rough JD keyword overlap:** {score:.1f}%")

            st.download_button(
                label="⬇️ Download optimized resume (.docx)",
                data=data,
                file_name="optimized_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

    except Exception as e:
        st.error(f"Something went wrong: {e}")
        st.stop()
