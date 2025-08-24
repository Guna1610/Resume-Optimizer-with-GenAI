# resume_optimizer.py
# One-file pipeline: parse resume/JD → Gemini JSON → replace SUMMARY, SKILLS, PROJECT EXPERIENCE in .docx (format preserved)

import os
import re
import json
from typing import Optional, List

from dotenv import load_dotenv
import google.generativeai as genai

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy
from docx.shared import Pt


# Fallback indents for SKILLS if no bullet template is found in the section
SKILLS_LEFT_INDENT_PT = 12     # overall indent
SKILLS_HANGING_INDENT_PT = 6   # hanging (so text aligns after bullet)

# Formatting knobs just for PROJECT EXPERIENCE
SECTION_HEADING_SPACE_AFTER_PT = 8        # space after "PROJECT EXPERIENCE"
PROJECT_TITLE_BOLD = True
PROJECT_TITLE_SPACE_AFTER_PT = 6          # space after each project title
PROJECT_BULLET_LEFT_INDENT_PT = 24        # overall bullet indent
PROJECT_BULLET_HANGING_PT = 12            # hanging indent so text lines up
PROJECT_BULLET_SPACE_AFTER_PT = 2         # space after each bullet line
PROJECT_BLOCK_SPACE_AFTER_PT = 8          # extra space after a project's bullet block
PROJECT_TITLE_SPACE_BEFORE_PT = 6    # space before each project title
PROJECT_TITLE_SPACE_AFTER_PT  = 4    # space after each project title
PROJECT_BLOCK_SPACE_AFTER_PT  = 12   # space after each project block (last bullet)

KNOWN_HEADINGS = {
    "SUMMARY",
    "SKILLS",
    "PROJECT EXPERIENCE",
    "WORK EXPERIENCE",
    "POTENTIAL PUBLICATIONS",
    "EDUCATION",
    "ACHIEVEMENTS",
    "EXTRA & CO-CURRICULAR ACTIVITIES",
}

# =========================
#  Config: API + Model
# =========================
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
MODEL_NAME = "gemini-1.5-flash"


# =========================
#  Utilities: Parsing
# =========================
def extract_resume_text(docx_path: str) -> str:
    """Read .docx and return text with newlines."""
    doc = Document(docx_path)
    lines: List[str] = []
    for p in doc.paragraphs:
        lines.append(p.text)
    return "\n".join(lines)


def extract_job_text(path: str) -> str:
    """Read a .txt JD. (Extend to PDF later if needed.)"""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


# =========================
#  Gemini: Strong ATS prompt
# =========================
def optimize_resume_sections(resume_text: str, job_text: str) -> dict:
    """Call Gemini to rewrite SUMMARY, SKILLS, PROJECT EXPERIENCE → JSON."""
    model = genai.GenerativeModel(MODEL_NAME)

    prompt = f"""
You are an advanced Resume Enhancement Assistant specialized in ATS (Applicant Tracking System) optimization.

Your task: Update and optimize the provided resume (.docx text) so that it achieves a 100% match with the provided Job Description (JD) in terms of ATS keyword alignment, relevance, and phrasing.

------------------
STRICT RULES:
------------------
1) Formatting & Structure
- The final Word document must keep the original formatting, section order, headings, bullet points, font styles, and layout.
- Do not add new sections or change section order.
- Do not modify Work Experience, Education, or Achievements.

2) Sections to Rewrite
- Rewrite ONLY these sections:
  - SUMMARY
  - SKILLS
  - PROJECT EXPERIENCE

3) Keyword & ATS Optimization
- Extract ALL critical keywords, skills, tools, and qualifications from the Job Description.
- Seamlessly integrate these keywords into SKILLS and PROJECT EXPERIENCE.
- Include exact technical keywords verbatim when appropriate (e.g., SQL, Python, Tableau, Azure, Databricks, DataOps, ETL, EHR, ML models, KPIs).

4) Content Enhancement
- Rewrite PROJECT EXPERIENCE bullets using strong action verbs + measurable outcomes.
- Emphasize business impact, scalability, and data-driven decision-making.
- Keep bullets concise (≤5 lines each) and ATS-friendly.

5) Intelligent Insertions (No Fabrication)
- If the JD emphasizes tools/skills logically aligned with the candidate’s background (e.g., MLOps, Azure Databricks, Healthcare Analytics, Cloud Data Pipelines), insert them naturally into SKILLS and PROJECT EXPERIENCE.
- Do NOT invent fake jobs, employers, or degrees. Only enhance existing projects and skills.

6) Output Format (CRITICAL)
- Return STRICT JSON ONLY with exactly these keys:
  - "skills": string 
              - Organize skills into 3–8 CATEGORY bullets (flexible). You MAY create new categories when appropriate.
              - Each bullet: start with "• ", then CATEGORY name, then a colon, then a comma-separated list of tools/skills.
              - Add ". " at the end of each category
              - If an item doesn’t fit an existing category, create a new appropriate category.
              - Keep the section organized into CATEGORY bullets 
              - Do NOT break every individual tool into its own bullet.
              
  - "projects": string 
              - Select the TOP 3 projects from the PROJECT LIBRARY that best match the job description.
              - Ignore irrelevant projects.
              - Exactly 3 projects unless fewer are available.
              - Each project must start with its TITLE in ALL CAPS (or Title Case).
              - After the title, all responsibilities/achievements must be written as bullet points.
              - Each bullet must begin with the "•" symbol followed by a tab or space.
              - Do not write responsibilities as plain paragraphs — only bulleted lists.
              - Example:
                     PROJECT TITLE
                        • First achievement with action verb, measurable outcome
                        • Second achievement with action verb, measurable outcome
- No extra commentary or markdown. JSON only.

------------------
RESUME:
{resume_text}

------------------
PROJECT LIBRARY (all projects you can choose from):
{project_library}

------------------
JOB DESCRIPTION:
{job_text}

OUTPUT JSON FORMAT:
{{
  "summary": "...",
  "skills": "...",
  "projects": "..."
}}
"""

    response = model.generate_content(
        prompt,
        generation_config={"response_mime_type": "application/json"}
    )

    try:
        data = json.loads(response.text)

        # normalize keys to be robust to variations like "project experience"
        norm = {k.strip().lower().replace("_", " "): v for k, v in data.items()}

        # map to canonical keys we use later
        result = {
            "summary": norm.get("summary", "").strip(),
            "skills": norm.get("skills", "").strip(),
            "projects": (
                    norm.get("projects")
                    or norm.get("project experience")
                    or norm.get("projectexperience")
                    or ""
            ).strip(),
        }

        # sanity check
        for k in ("skills", "projects"):
            if not result[k]:
                raise KeyError(f"Missing or empty key after normalization: {k}")

        return result

    except Exception:
        print("⚠️ Gemini returned non-JSON or malformed/variant-key output. Raw response below:")
        print(response.text)
        raise


# ---------- FONT UTILS ----------
def _force_times_new_roman(paragraph: Paragraph):
    """Force Times New Roman, 12pt for all runs in a paragraph."""
    for r in paragraph.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)  # <-- force size 12pt
        # ensure Word respects it for all scripts
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:eastAsia"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")


# =========================
#  DOCX Utilities: Formatting-preserving insertion
# =========================

def extract_project_library(path: str) -> str:
    """Read a text/markdown file with all projects listed."""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def _is_heading(p: Paragraph) -> bool:
    """Treat as heading only if it matches known headings (by text) or uses a Word heading style."""
    name = (p.style.name or "").lower() if p.style is not None else ""
    txt = p.text.strip()
    # If the paragraph uses a Word heading style, accept it
    if name.startswith("heading"):
        return True
    # Normalize spaces for robust matching
    txt_norm = " ".join(txt.split()).upper()
    return txt_norm in KNOWN_HEADINGS

def _apply_bullet_alignment_from_template(template_para: Paragraph, dst_para: Paragraph):
    """Copy bullet paragraph spacing/indents from a template paragraph."""
    if template_para is None:
        return
    _apply_para_style_and_format(template_para, dst_para)  # copies style + paragraph_format

def _apply_skills_fallback_indents(p: Paragraph):
    """Clean, consistent fallback indentation for Skills bullets when no template exists."""
    pf = p.paragraph_format
    pf.left_indent = Pt(SKILLS_LEFT_INDENT_PT)
    pf.first_line_indent = Pt(-SKILLS_HANGING_INDENT_PT)

def _apply_project_bullet_indents(p):
    pf = p.paragraph_format
    pf.left_indent = Pt(PROJECT_BULLET_LEFT_INDENT_PT)
    pf.first_line_indent = Pt(-PROJECT_BULLET_HANGING_PT)
    pf.space_after = Pt(PROJECT_BULLET_SPACE_AFTER_PT)

def _bold_entire_paragraph(p, text):
    # clear existing runs, add one bold run
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    run = p.add_run(text)
    run.bold = True

def _is_section_heading(p: Paragraph) -> bool:
    name = (p.style.name or "").lower() if p.style is not None else ""
    if name.startswith("heading"):
        return True
    txt_norm = " ".join(p.text.strip().split()).upper()
    return txt_norm in KNOWN_HEADINGS

def _paragraph_has_bullets(p: Paragraph) -> bool:
    pPr = p._p.pPr
    return (pPr is not None and pPr.numPr is not None)

def _find_section_bounds(doc, heading_text):
    paragraphs = doc.paragraphs
    target = " ".join(heading_text.split()).upper()
    start_idx = None
    for i, p in enumerate(paragraphs):
        txt_norm = " ".join(p.text.strip().split()).upper()
        if txt_norm == target:
            start_idx = i
            break
    if start_idx is None:
        return None
    end_idx = len(paragraphs)
    for j in range(start_idx + 1, len(paragraphs)):
        if _is_section_heading(paragraphs[j]):
            end_idx = j
            break
    return start_idx, end_idx

    # Find the next real section heading
    end_idx = len(paragraphs)
    for j in range(start_idx + 1, len(paragraphs)):
        if _is_heading(paragraphs[j]):
            end_idx = j
            break

    return (start_idx, end_idx)

def _add_paragraph_after(para: Paragraph) -> Paragraph:
    new_p = OxmlElement('w:p')
    para._element.addnext(new_p)
    return Paragraph(new_p, para._parent)


def _copy_numPr(src_para: Paragraph, dst_para: Paragraph) -> None:
    # copy Word bullet/numbering definition
    src_pPr = src_para._p.pPr
    if src_pPr is None or src_pPr.numPr is None:
        return
    dst_pPr = dst_para._p.get_or_add_pPr()
    for child in list(dst_pPr):
        if child.tag == qn('w:numPr'):
            dst_pPr.remove(child)
    dst_pPr.append(deepcopy(src_pPr.numPr))

def _looks_like_title(line: str) -> bool:
    t = line.strip()
    if len(t) < 4:
        return False
    letters = [ch for ch in t if ch.isalpha()]
    if not letters:
        return False
    upper_ratio = sum(ch.isupper() for ch in letters) / len(letters)
    # treat mostly-upper lines as titles (e.g., your project headings)
    return upper_ratio > 0.65 and not t.startswith(("•", "-", "–", "*"))

def _apply_para_style_and_format(src_para: Paragraph, dst_para: Paragraph) -> None:
    # copy style & spacing/indents exactly from src to dst
    if src_para.style is not None:
        dst_para.style = src_para.style
    sp_src = src_para.paragraph_format
    sp_dst = dst_para.paragraph_format
    sp_dst.left_indent = sp_src.left_indent
    sp_dst.first_line_indent = sp_src.first_line_indent
    sp_dst.space_before = sp_src.space_before
    sp_dst.space_after = sp_src.space_after
    sp_dst.line_spacing = sp_src.line_spacing
    sp_dst.line_spacing_rule = sp_src.line_spacing_rule
    sp_dst.keep_together = sp_src.keep_together
    sp_dst.keep_with_next = sp_src.keep_with_next
    sp_dst.widow_control = sp_src.widow_control


def _clean_leading_bullet(text: str) -> str:
    s = text.strip()
    for b in ("• ", "- ", "– ", "* "):
        if s.startswith(b):
            return s[len(b):].strip()
    return s

def replace_section(doc: Document, section_name: str, new_content: str) -> None:
    bounds = _find_section_bounds(doc, section_name)
    if not bounds:
        print(f"⚠️ Section '{section_name}' not found. Skipping.")
        return

    start_idx, end_idx = bounds
    paragraphs = doc.paragraphs
    heading_para = paragraphs[start_idx]

    # ---- discover templates BEFORE deletion (for SKILLS alignment) ----
    skills_bullet_template = None
    skills_text_template = None
    if section_name.strip().upper() == "SKILLS":
        for k in range(start_idx + 1, end_idx):
            p = paragraphs[k]
            if (p.text or "").strip():
                if _paragraph_has_bullets(p) and skills_bullet_template is None:
                    skills_bullet_template = p
                if skills_text_template is None:
                    skills_text_template = p  # first body para as general text style
            if skills_bullet_template is not None and skills_text_template is not None:
                break
        if skills_text_template is None:
            skills_text_template = heading_para


    # ---- discover templates BEFORE deletion (for PROJECT EXPERIENCE alignment) ----
    title_template = None
    bullet_template = None
    if section_name.strip().upper() == "PROJECT EXPERIENCE":
        for k in range(start_idx + 1, end_idx):
            p = paragraphs[k]
            if (p.text or "").strip():
                if _paragraph_has_bullets(p):
                    if bullet_template is None:
                        bullet_template = p
                else:
                    if title_template is None:
                        title_template = p
            if title_template is not None and bullet_template is not None:
                break

    # Fallback templates if section had no body yet
    if title_template is None:
        title_template = heading_para
    if bullet_template is None:
        # if no bullets found in section, try to use first body para as style
        bullet_template = title_template

    # Delete old content (keep heading)
    for idx in range(end_idx - 1, start_idx, -1):
        p_elm = paragraphs[idx]._element
        p_elm.getparent().remove(p_elm)

    # Re-acquire after deletion
    paragraphs = doc.paragraphs
    heading_para = paragraphs[start_idx]

    # Choose a general text template for SUMMARY/SKILLS (first body para or heading)
    if section_name.strip().upper() != "PROJECT EXPERIENCE":
        text_template = paragraphs[start_idx + 1] if (start_idx + 1) < len(paragraphs) and (start_idx + 1) < end_idx else heading_para

    lines = [ln for ln in new_content.splitlines() if ln.strip() != ""]
    anchor = heading_para
    in_project_block = False

    for raw in lines:
        line = raw.rstrip()
        is_title = _looks_like_title(line)

        new_para = _add_paragraph_after(anchor)

        if section_name.strip().upper() == "PROJECT EXPERIENCE":
            if is_title:
                # copy EXACT alignment/style from title_template
                _apply_para_style_and_format(title_template, new_para)

                # titles are bold but keep same left indent / spacing as template
                for r in list(new_para.runs):
                    r._element.getparent().remove(r._element)
                run = new_para.add_run(line.strip())
                run.bold = True
                _force_times_new_roman(new_para)

                # spacing around title
                pf = new_para.paragraph_format
                pf.space_before = Pt(PROJECT_TITLE_SPACE_BEFORE_PT)
                pf.space_after = Pt(PROJECT_TITLE_SPACE_AFTER_PT)

                in_project_block = False

            else:
                # bullet: apply bullet template style & numbering exactly
                _apply_para_style_and_format(bullet_template, new_para)
                _copy_numPr(bullet_template, new_para)

                # write text
                clean_text = _clean_leading_bullet(line)
                for r in list(new_para.runs):
                    r._element.getparent().remove(r._element)
                new_para.add_run(clean_text)
                _force_times_new_roman(new_para)

                # default bullet spacing
                pf = new_para.paragraph_format
                pf.space_after = Pt(2)

                in_project_block = True

        else:
            # SUMMARY / SKILLS path
            if section_name.strip().upper() == "SKILLS":
                # Always create bullets for each line
                # Prefer SKILLS bullet template if found
                template_for_bullets = skills_bullet_template or skills_text_template or heading_para

                # Start paragraph from the text template for consistent font/spacing
                _apply_para_style_and_format(skills_text_template or heading_para, new_para)

                # Copy numbering (bullets). If text template has no bullets, try any bullet in doc.
                if _paragraph_has_bullets(template_for_bullets):
                    _copy_numPr(template_for_bullets, new_para)
                    # Copy the exact bullet alignment from template
                    _apply_bullet_alignment_from_template(template_for_bullets, new_para)
                else:
                    # Fallback: try to borrow numPr from any bullet in the doc
                    borrowed = False
                    for p_any in doc.paragraphs:
                        if _paragraph_has_bullets(p_any):
                            _copy_numPr(p_any, new_para)
                            _apply_bullet_alignment_from_template(p_any, new_para)
                            borrowed = True
                            break
                    if not borrowed:
                        # Last resort: apply neat fallback indents so bullets don't drift
                        _apply_skills_fallback_indents(new_para)

                # Clear runs and write (with bold category before colon)
                for r in list(new_para.runs):
                    r._element.getparent().remove(r._element)

                text = _clean_leading_bullet(line)
                if ":" in text:
                    category, rest = text.split(":", 1)
                    run_cat = new_para.add_run(category.strip() + ": ")
                    run_cat.bold = True
                    run_cat.font.name = "Times New Roman"

                    run_rest = new_para.add_run(rest.strip())
                    run_rest.bold = False
                    run_rest.font.name = "Times New Roman"
                else:
                    run = new_para.add_run(text.strip())
                    run.font.name = "Times New Roman"

                _force_times_new_roman(new_para)

            else:
                # SUMMARY (unchanged except TNR + copy bullets only if present)
                _apply_para_style_and_format(text_template, new_para)
                if _paragraph_has_bullets(text_template) or line.strip().startswith(("•", "-", "–", "*")):
                    _copy_numPr(text_template, new_para)
                for r in list(new_para.runs):
                    r._element.getparent().remove(r._element)
                new_para.add_run(_clean_leading_bullet(line))
                _force_times_new_roman(new_para)

        anchor = new_para

    # Optional: add a little space after the last bullet block of PROJECT EXPERIENCE
    # Final project block (ensure spacing after last project)
    if section_name.strip().upper() == "PROJECT EXPERIENCE" and in_project_block:
        pf_last = anchor.paragraph_format
        pf_last.space_after = Pt(PROJECT_BLOCK_SPACE_AFTER_PT)


# =========================
#  Orchestrator
# =========================
def optimize_resume(resume_path: str, job_path: str, output_path: str = "optimized_resume.docx") -> None:
    # Read inputs
    resume_text = extract_resume_text(resume_path)
    job_text = extract_job_text(job_path)

    # Gemini → JSON
    project_library = extract_project_library("projects.txt")
    optimized = optimize_resume_sections(resume_text, job_text, project_library)

    # Load original docx
    doc = Document(resume_path)

    # Replace ONLY these three sections
    replace_section(doc, "SUMMARY", optimized["summary"])
    replace_section(doc, "SKILLS", optimized["skills"])
    replace_section(doc, "PROJECT EXPERIENCE", optimized["projects"])

    # Save
    doc.save(output_path)
    print(f"✅ Optimized resume saved as: {output_path}")


# =========================
#  Run
# =========================
if __name__ == "__main__":
    # Adjust filenames if needed
    optimize_resume(resume_path="sample_resume.docx",
                    job_path="sample_job.txt",
                    output_path="optimized_resume.docx")
