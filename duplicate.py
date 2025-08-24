# resume_optimizer.py
# End-to-end: read resume/JD + PROJECT LIBRARY → Gemini JSON → replace SUMMARY, SKILLS, PROJECT EXPERIENCE in .docx (format preserved)
# - Project Experience: picks TOP 3 relevant projects from projects.txt
# - Project Experience font forced to Times New Roman 12pt
# - Skills bullets + bold categories; formatting preserved

import os
import re
import json
from typing import Optional, List
from copy import deepcopy

from dotenv import load_dotenv
import google.generativeai as genai

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


# ===== User-tunable formatting =====
# Fallback indents for SKILLS if no bullet template is found in the section
SKILLS_LEFT_INDENT_PT = 12
SKILLS_HANGING_INDENT_PT = 6

# Spacing around Project blocks
PROJECT_TITLE_SPACE_BEFORE_PT = 6    # before each project title
PROJECT_TITLE_SPACE_AFTER_PT  = 4    # after each project title
PROJECT_BULLET_SPACE_AFTER_PT = 2    # after each bullet
PROJECT_BLOCK_SPACE_AFTER_PT  = 12   # after each project block (last bullet)

KNOWN_HEADINGS = {
    "SUMMARY",
    "SKILLS",
    "PROJECT EXPERIENCE",
    "WORK EXPERIENCE",
    "POTENTIAL PUBLICATIONS",
    "EDUCATION",
    "ACHIEVEMENTS",
    "EXTRA & CO-CURRICULAR ACTIVITIES",
}

# ===== API / Model =====
load_dotenv()
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
MODEL_NAME = "gemini-1.5-flash"


# ===== IO helpers =====
def extract_resume_text(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_job_text(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def extract_project_library(path: str) -> str:
    """
    Read a text/markdown file listing ALL projects you can showcase.
    Format: blocks like
    PROJECT TITLE
    • bullet
    • bullet

    <blank line>
    NEXT PROJECT
    • bullet ...
    """
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


# ===== Gemini: prompt + post-process =====
def _is_project_title_line(line: str) -> bool:
    t = line.strip()
    if not t or t.startswith(("•", "-", "–", "*")):
        return False
    letters = [ch for ch in t if ch.isalpha()]
    if not letters:
        return False
    # treat UPPER-ish or Title Case as title
    upper_ratio = sum(ch.isupper() for ch in letters) / len(letters)
    return upper_ratio > 0.65 or t.istitle()

def _split_projects_into_blocks(projects_text: str) -> list[list[str]]:
    """Split the projects string into blocks: [ [title, bullet, ...], [title, ...], ... ]"""
    lines = [ln.rstrip() for ln in projects_text.splitlines()]
    blocks = []
    cur = []
    for ln in lines:
        if not ln.strip():
            if cur and cur[-1].strip():
                cur.append("")  # keep a single blank inside block if needed
            continue
        if _is_project_title_line(ln):
            if cur:
                # trim trailing blanks
                while cur and not cur[-1].strip():
                    cur.pop()
                blocks.append(cur)
            cur = [ln]
        else:
            if not cur:
                cur = [ln]  # bullets before any title → still keep
            else:
                cur.append(ln)
    if cur:
        while cur and not cur[-1].strip():
            cur.pop()
        blocks.append(cur)
    return blocks

def _keep_top_n_projects(projects_text: str, n: int = 3) -> str:
    blocks = _split_projects_into_blocks(projects_text)
    top = blocks[:n]
    return "\n\n".join("\n".join(b) for b in top)

def optimize_resume_sections(resume_text: str, job_text: str, project_library: str) -> dict:
    """
    Ask Gemini for STRICT JSON with keys:
      - "summary": string
      - "skills": string
      - "projects": string  (TOP 3 projects selected from project library)
    """
    model = genai.GenerativeModel(MODEL_NAME)

    prompt = f"""
You are an advanced Resume Enhancement Assistant specialized in ATS (Applicant Tracking System) optimization.

Your task: Update and optimize the provided resume (.docx text) so that it achieves a 100% match with the provided Job Description (JD) in terms of ATS keyword alignment, relevance, and phrasing.

------------------
STRICT RULES:
------------------
1) Formatting & Structure
- The final Word document must keep the original formatting, section order, headings, bullet points, font styles, and layout.
- Do not add new sections or change section order.
- Do not modify Work Experience, Education, or Achievements.

2) Sections to Rewrite
- Rewrite ONLY these sections:
  - SUMMARY
  - SKILLS
  - PROJECT EXPERIENCE

3) Keyword & ATS Optimization
- Extract ALL critical keywords, skills, tools, and qualifications from the Job Description.
- Seamlessly integrate these keywords into SKILLS and PROJECT EXPERIENCE.
- Include exact technical keywords verbatim when appropriate (e.g., SQL, Python, Tableau, Azure, Databricks, DataOps, ETL, EHR, ML models, KPIs).

4) Content Enhancement
- Rewrite PROJECT EXPERIENCE bullets using strong action verbs + measurable outcomes.
- Emphasize business impact, scalability, and data-driven decision-making.
- Keep bullets concise (≤5 lines each) and ATS-friendly.

5) Intelligent Insertions (No Fabrication)
- If the JD emphasizes tools/skills logically aligned with the candidate’s background (e.g., MLOps, Azure Databricks, Healthcare Analytics, Cloud Data Pipelines), insert them naturally into SKILLS and PROJECT EXPERIENCE.
- Do NOT invent fake jobs, employers, or degrees. Only enhance existing projects and skills.

6) Output Format (CRITICAL)
- Return STRICT JSON ONLY with exactly these keys:
  - "skills": string 
              - Organize skills into 3–8 CATEGORY bullets (flexible). You MAY create new categories when appropriate.
              - Each bullet: start with "• ", then CATEGORY name, then a colon, then a comma-separated list of tools/skills.
              - Add ". " at the end of each category
              - If an item doesn’t fit an existing category, create a new appropriate category.
              - Keep the section organized into CATEGORY bullets 
              - Do NOT break every individual tool into its own bullet.
              
  - "projects": string 
              - Select the TOP 3 projects from the PROJECT LIBRARY that best match the job description.
              - Ignore irrelevant projects.
              - Exactly 3 projects unless fewer are available.
              - Each project must start with its TITLE in ALL CAPS (or Title Case).
              - After the title, all responsibilities/achievements must be written as bullet points.
              - Each bullet must begin with the "•" symbol followed by a tab or space.
              - Do not write responsibilities as plain paragraphs — only bulleted lists.
              - Example:
                     PROJECT TITLE
                        • First achievement with action verb, measurable outcome
                        • Second achievement with action verb, measurable outcome
- No extra commentary or markdown. JSON only.

------------------
RESUME:
{resume_text}

------------------
PROJECT LIBRARY (all projects you can choose from):
{project_library}

------------------
JOB DESCRIPTION:
{job_text}

OUTPUT JSON FORMAT:
{{
  "summary": "...",
  "skills": "...",
  "projects": "..."
}}
"""

    response = model.generate_content(
        prompt,
        generation_config={"response_mime_type": "application/json"}
    )

    try:
        data = json.loads(response.text)
        norm = {k.strip().lower().replace("_", " "): v for k, v in data.items()}

        summary  = (norm.get("summary")  or "").strip()
        skills   = (norm.get("skills")   or "").strip()
        projects = (norm.get("projects") or norm.get("project experience") or "").strip()

        # Safety: enforce top-3 blocks even if model sends more
        if projects:
            projects = _keep_top_n_projects(projects, n=3)

        # Minimal sanity
        for k, v in [("skills", skills), ("projects", projects)]:
            if not v:
                raise KeyError(f"Missing or empty key: {k}")

        return {"summary": summary, "skills": skills, "projects": projects}

    except Exception:
        print("⚠️ Gemini returned non-JSON or malformed output. Raw response below:")
        print(response.text)
        raise


# ===== DOCX formatting utils =====
def _is_section_heading(p: Paragraph) -> bool:
    name = (p.style.name or "").lower() if p.style is not None else ""
    if name.startswith("heading"):
        return True
    txt_norm = " ".join(p.text.strip().split()).upper()
    return txt_norm in KNOWN_HEADINGS

def _find_section_bounds(doc: Document, heading_text: str):
    paragraphs = doc.paragraphs
    target = " ".join(heading_text.split()).upper()
    start_idx = None
    for i, p in enumerate(paragraphs):
        if " ".join(p.text.strip().split()).upper() == target:
            start_idx = i
            break
    if start_idx is None:
        return None
    end_idx = len(paragraphs)
    for j in range(start_idx + 1, len(paragraphs)):
        if _is_section_heading(paragraphs[j]):
            end_idx = j
            break
    return start_idx, end_idx

def _paragraph_has_bullets(p: Paragraph) -> bool:
    pPr = p._p.pPr
    return (pPr is not None and pPr.numPr is not None)

def _apply_para_style_and_format(src_para: Paragraph, dst_para: Paragraph) -> None:
    if src_para.style is not None:
        dst_para.style = src_para.style
    sp_src = src_para.paragraph_format
    sp_dst = dst_para.paragraph_format
    sp_dst.left_indent = sp_src.left_indent
    sp_dst.first_line_indent = sp_src.first_line_indent
    sp_dst.space_before = sp_src.space_before
    sp_dst.space_after = sp_src.space_after
    sp_dst.line_spacing = sp_src.line_spacing
    sp_dst.line_spacing_rule = sp_src.line_spacing_rule
    sp_dst.keep_together = sp_src.keep_together
    sp_dst.keep_with_next = sp_src.keep_with_next
    sp_dst.widow_control = sp_src.widow_control

def _copy_numPr(src_para: Paragraph, dst_para: Paragraph) -> None:
    src_pPr = src_para._p.pPr
    if src_pPr is None or src_pPr.numPr is None:
        return
    dst_pPr = dst_para._p.get_or_add_pPr()
    for child in list(dst_pPr):
        if child.tag == qn('w:numPr'):
            dst_pPr.remove(child)
    dst_pPr.append(deepcopy(src_pPr.numPr))

def _add_paragraph_after(para: Paragraph) -> Paragraph:
    new_p = OxmlElement('w:p')
    para._element.addnext(new_p)
    return Paragraph(new_p, para._parent)

def _clean_leading_bullet(text: str) -> str:
    """
    Remove any leading bullet glyphs/dashes/tabs/spaces so we don't get a visual 'double bullet'
    when the paragraph itself already has Word bullets applied.
    """
    s = text.lstrip()  # strip leading spaces
    # common bullet-like prefixes
    prefixes = [
        "•", "●", "◦", "▪", "‣", "·",  # bullet glyphs
        "-", "–", "—",                 # dashes
        "*",                           # asterisk
    ]
    # strip one prefix + optional following spaces/tabs
    for pref in prefixes:
        if s.startswith(pref):
            s = s[len(pref):].lstrip(" \t")
            break
    # also handle cases like "•\t", "-\t", "*\t"
    if s.startswith(("\t", " ")):
        s = s.lstrip(" \t")
    return s
def _looks_like_title(line: str) -> bool:
    t = line.strip()
    if len(t) < 4:
        return False
    letters = [ch for ch in t if ch.isalpha()]
    if not letters:
        return False
    upper_ratio = sum(ch.isupper() for ch in letters) / len(letters)
    return upper_ratio > 0.65 and not t.startswith(("•", "-", "–", "*"))

def _force_tnr_12(paragraph: Paragraph) -> None:
    """Force Times New Roman 12pt for all runs in a paragraph."""
    for r in paragraph.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:eastAsia"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")

def _force_tnr(paragraph: Paragraph) -> None:
    """Force Times New Roman (keep whatever size the paragraph style gives)."""
    for r in paragraph.runs:
        r.font.name = "Times New Roman"
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:eastAsia"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")


# ===== Section replacer =====
def replace_section(doc: Document, section_name: str, new_content: str) -> None:
    bounds = _find_section_bounds(doc, section_name)
    if not bounds:
        print(f"⚠️ Section '{section_name}' not found. Skipping.")
        return

    # ---- spacing knobs ----
    HEADING_SPACE_AFTER_PT = 12  # after "PROJECT EXPERIENCE"

    start_idx, end_idx = bounds
    paragraphs = doc.paragraphs
    heading_para = paragraphs[start_idx]

    # ---- discover templates BEFORE deletion ----
    skills_bullet_template = None
    skills_text_template = None
    title_template = None
    bullet_template = None

    if section_name.strip().upper() == "SKILLS":
        for k in range(start_idx + 1, end_idx):
            p = paragraphs[k]
            if (p.text or "").strip():
                if _paragraph_has_bullets(p) and skills_bullet_template is None:
                    skills_bullet_template = p
                if skills_text_template is None:
                    skills_text_template = p
                if skills_bullet_template and skills_text_template:
                    break
        if skills_text_template is None:
            skills_text_template = heading_para

    if section_name.strip().upper() == "PROJECT EXPERIENCE":
        for k in range(start_idx + 1, end_idx):
            p = paragraphs[k]
            if (p.text or "").strip():
                if _paragraph_has_bullets(p) and bullet_template is None:
                    bullet_template = p
                if (not _paragraph_has_bullets(p)) and title_template is None:
                    title_template = p
                if title_template and bullet_template:
                    break
        if title_template is None:
            title_template = heading_para
        if bullet_template is None:
            bullet_template = title_template

    # ---- delete old content (keep heading) ----
    for idx in range(end_idx - 1, start_idx, -1):
        elm = paragraphs[idx]._element
        elm.getparent().remove(elm)

    # ---- re-acquire after deletion ----
    paragraphs = doc.paragraphs
    heading_para = paragraphs[start_idx]

    # add space after the PROJECT EXPERIENCE heading
    if section_name.strip().upper() == "PROJECT EXPERIENCE":
        heading_para.paragraph_format.space_after = Pt(HEADING_SPACE_AFTER_PT)

    # general text template for SUMMARY/SKILLS
    if section_name.strip().upper() != "PROJECT EXPERIENCE":
        text_template = paragraphs[start_idx + 1] if (start_idx + 1) < len(paragraphs) and (start_idx + 1) < end_idx else heading_para

    lines = [ln for ln in new_content.splitlines() if ln.strip()]
    anchor = heading_para

    # tracking for project block spacing
    in_project_block = False
    first_project = True
    last_bullet_para: Paragraph | None = None

    for i, raw in enumerate(lines):
        line = raw.rstrip()
        is_title = _looks_like_title(line)
        new_para = _add_paragraph_after(anchor)

        if section_name.strip().upper() == "PROJECT EXPERIENCE":
            if is_title:
                # close previous block with extra space (apply to the LAST BULLET paragraph)
                if in_project_block and last_bullet_para is not None:
                    last_bullet_para.paragraph_format.space_after = Pt(PROJECT_BLOCK_SPACE_AFTER_PT)
                    in_project_block = False
                    last_bullet_para = None

                # Title: bold, spacing, TNR 12
                _apply_para_style_and_format(title_template, new_para)
                for r in list(new_para.runs):
                    r._element.getparent().remove(r._element)
                run = new_para.add_run(line.strip())
                run.bold = True
                _force_tnr_12(new_para)

                pf = new_para.paragraph_format
                pf.space_before = Pt(0 if first_project else PROJECT_TITLE_SPACE_BEFORE_PT)
                pf.space_after  = Pt(PROJECT_TITLE_SPACE_AFTER_PT)
                first_project = False

            else:
                # Bullet: copy bullet template style + numbering, TNR 12
                _apply_para_style_and_format(bullet_template, new_para)
                _copy_numPr(bullet_template, new_para)

                # write clean text (no leading bullet glyph)
                for r in list(new_para.runs):
                    r._element.getparent().remove(r._element)
                new_para.add_run(_clean_leading_bullet(line))
                _force_tnr_12(new_para)

                # per-bullet spacing
                new_para.paragraph_format.space_after = Pt(PROJECT_BULLET_SPACE_AFTER_PT)
                in_project_block = True
                last_bullet_para = new_para

        elif section_name.strip().upper() == "SKILLS":
            # Skills bullets with category bold before colon
            template_for_bullets = skills_bullet_template or skills_text_template or heading_para
            _apply_para_style_and_format(skills_text_template or heading_para, new_para)
            if _paragraph_has_bullets(template_for_bullets):
                _copy_numPr(template_for_bullets, new_para)
                _apply_para_style_and_format(template_for_bullets, new_para)
            else:
                # fallback: borrow any bullet; else set a neat indent
                borrowed = False
                for p_any in doc.paragraphs:
                    if _paragraph_has_bullets(p_any):
                        _copy_numPr(p_any, new_para)
                        _apply_para_style_and_format(p_any, new_para)
                        borrowed = True
                        break
                if not borrowed:
                    pf = new_para.paragraph_format
                    pf.left_indent = Pt(SKILLS_LEFT_INDENT_PT)
                    pf.first_line_indent = Pt(-SKILLS_HANGING_INDENT_PT)

            for r in list(new_para.runs):
                r._element.getparent().remove(r._element)
            text = _clean_leading_bullet(line)
            if ":" in text:
                cat, rest = text.split(":", 1)
                rc = new_para.add_run(cat.strip() + ": ")
                rc.bold = True
                rr = new_para.add_run(rest.strip())
            else:
                new_para.add_run(text.strip())
            _force_tnr(new_para)  # keep paragraph style size for Skills

        else:
            # SUMMARY
            _apply_para_style_and_format(text_template, new_para)
            for r in list(new_para.runs):
                r._element.getparent().remove(r._element)
            new_para.add_run(_clean_leading_bullet(line))
            _force_tnr(new_para)

        # if this is the last line of the entire PROJECT EXPERIENCE input and we're in a block,
        # ensure block spacing after the last bullet
        if section_name.strip().upper() == "PROJECT EXPERIENCE":
            is_last_line = (i == len(lines) - 1)
            if is_last_line and in_project_block and last_bullet_para is not None:
                last_bullet_para.paragraph_format.space_after = Pt(PROJECT_BLOCK_SPACE_AFTER_PT)

        anchor = new_para


# ===== Orchestrator =====
def optimize_resume(resume_path: str, job_path: str, project_library_path: str, output_path: str = "optimized_resume.docx") -> None:
    # Read inputs
    resume_text = extract_resume_text(resume_path)
    job_text = extract_job_text(job_path)
    project_library = extract_project_library(project_library_path)

    # Gemini → JSON
    optimized = optimize_resume_sections(resume_text, job_text, project_library)

    # Load original docx
    doc = Document(resume_path)

    # Replace SUMMARY, SKILLS, PROJECT EXPERIENCE
    # (If you want to remove SUMMARY entirely, comment out this line and also delete the section from doc beforehand.)
    replace_section(doc, "SUMMARY", optimized["summary"])
    replace_section(doc, "SKILLS", optimized["skills"])
    replace_section(doc, "PROJECT EXPERIENCE", optimized["projects"])

    # Save
    doc.save(output_path)
    print(f"✅ Optimized resume saved as: {output_path}")


# ===== Run =====
if __name__ == "__main__":
    optimize_resume(
        resume_path="sample_resume.docx",            # your original resume file
        job_path="sample_job.txt",            # JD text file
        project_library_path="projects.txt",  # NEW: your full project library
        output_path="optimized_resume.docx"
    )
